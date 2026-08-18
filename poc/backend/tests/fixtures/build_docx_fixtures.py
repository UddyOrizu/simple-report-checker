"""One-off generator for the .docx test fixtures. Run with `python build_docx_fixtures.py`."""

import os

from docx import Document

FIXTURES_DIR = os.path.dirname(__file__)


def build_sample_report():
    doc = Document()
    doc.add_heading("Q3 Business Report", level=0)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report summarizes our performance for the third quarter, covering financial "
        "results, regional performance, and forward-looking guidance."
    )

    doc.add_heading("Financial Highlights", level=1)
    doc.add_paragraph(
        "Revenue grew 12% YoY, driven primarily by APAC expansion and improved cost "
        "discipline, positioning us ahead of our closest competitor."
    )
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Metric", "Value"),
        ("Revenue (current period)", "$112M"),
        ("Revenue (prior period)", "$100M"),
    ]
    for row_idx, (label, value) in enumerate(rows):
        table.cell(row_idx, 0).text = label
        table.cell(row_idx, 1).text = value

    doc.add_heading("Outlook", level=1)
    doc.add_paragraph(
        "We expect continued momentum into the fourth quarter, subject to macroeconomic conditions."
    )

    doc.save(os.path.join(FIXTURES_DIR, "sample_report.docx"))


def build_short_memo():
    doc = Document()
    doc.add_heading("Internal Memo", level=1)
    doc.add_paragraph("Team,")
    doc.add_paragraph(
        "Quick update: the vendor contract renewal is on track for next month. No action "
        "needed from anyone outside procurement at this time."
    )
    doc.add_paragraph("Thanks,\nOps Team")
    doc.save(os.path.join(FIXTURES_DIR, "short_memo.docx"))


if __name__ == "__main__":
    build_sample_report()
    build_short_memo()
    print("wrote sample_report.docx and short_memo.docx")
